from io import BytesIO

from docx import Document
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph


# TXT
def create_txt(text: str) -> bytes:
    """
    Creates a TXT file.
    """

    return text.encode("utf-8")


# DOCX
def create_docx(text: str) -> BytesIO:
    """
    Creates a DOCX file.
    """

    document = Document()

    document.add_heading(
        "OmniSummarizer AI",
        level=1,
    )

    document.add_paragraph(text)

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return buffer


# PDF
def create_pdf(text: str) -> BytesIO:
    """
    Creates a PDF file.
    """

    buffer = BytesIO()

    document = SimpleDocTemplate(buffer)

    styles = getSampleStyleSheet()

    story = [
        Paragraph(
            "<b>OmniSummarizer AI</b>",
            styles["Heading1"],
        ),
        Paragraph(text, styles["BodyText"]),
    ]

    document.build(story)

    buffer.seek(0)

    return buffer